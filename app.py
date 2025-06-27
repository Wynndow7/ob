import streamlit as st
from docx import Document
import openai
import json
import tempfile
import os

st.set_page_config(page_title="Course Builder", layout="wide")
st.title("📘 GPT-Powered Collaborative Course Builder")

# --- Upload Pre-Outline ---
st.header("Step 1: Upload Pre-Outline Document")
pre_outline_doc = st.file_uploader("Upload .docx Pre-Outline", type="docx")

parsed_text = ""
if pre_outline_doc:
    doc = Document(pre_outline_doc)
    parsed_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    st.success("Pre-outline loaded successfully.")
    st.text_area("Preview Extracted Pre-Outline Text", parsed_text, height=200)

# --- GPT Outcome Analyzer ---
st.header("Step 2: Generate and Edit Outcome Analysis")
if parsed_text:
    openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")
    if openai_api_key:
        if st.button("🔍 Analyze Learning Outcomes with GPT"):
            with st.spinner("Analyzing via GPT..."):
                prompt = f"You are an instructional designer. Based on the following pre-outline, answer questions about the course learning outcomes:\n\n{parsed_text}\n\nGenerate thoughtful, practical responses for instructional planning."
                try:
                    openai.api_key = openai_api_key
                    response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7
                    )
                    gpt_response = response['choices'][0]['message']['content']
                    st.session_state['gpt_answers'] = gpt_response
                except Exception as e:
                    st.error(f"Error from OpenAI: {e}")

if 'gpt_answers' in st.session_state:
    st.subheader("✏️ SME + Designer Edits")
    edited_answers = st.text_area("Edit GPT's Suggested Outcomes Answers", st.session_state['gpt_answers'], height=300)

    if st.button("✅ Finalize and Build Outline"):
        st.session_state['final_outline_input'] = edited_answers

# --- Generate Outline (Placeholder for GPT) ---
if 'final_outline_input' in st.session_state:
    st.header("Step 3: Generated Course Outline (Editable)")
    outline_text = f"# Course Outline\n\n## Module 1\n- Topic: Introduction\n- Learning Outcome: {st.session_state['final_outline_input'][:100]}..."
    edited_outline = st.text_area("Edit the Outline Below", outline_text, height=300)

    if st.button("📝 Generate Final Course Manuscript"):
        st.session_state['final_manuscript_input'] = edited_outline

# --- Manuscript Builder ---
def generate_docx(text):
    doc = Document()
    doc.add_heading("Generated Course Manuscript", level=1)
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "course_manuscript.docx")
    doc.save(file_path)
    return file_path

if 'final_manuscript_input' in st.session_state:
    st.header("Step 4: Final Course Manuscript Preview")
    manuscript_text = st.session_state['final_manuscript_input'] + "\n\n(Full manuscript would be expanded here.)"
    st.text_area("Review and Export Manuscript", manuscript_text, height=300)

    docx_file_path = generate_docx(manuscript_text)
    with open(docx_file_path, "rb") as file:
        st.download_button(
            label="⬇️ Download as Word Document (.docx)",
            data=file,
            file_name="course_manuscript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
